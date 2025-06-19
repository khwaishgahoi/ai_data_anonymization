<<<<<<< HEAD
import os
import pandas as pd
import random
import string
import re
from flask import Flask, request, send_file, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from docx import Document
import fitz  # PyMuPDF
import pdfplumber
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ========== Helper Functions ==========

def randomize_string(s):
    return ''.join(random.choice(string.ascii_letters) if c.isalpha() else c for c in s)

def randomize_phone(phone):
    phone = list(phone)
    for i in random.sample(range(min(5, len(phone))), k=3):
        if phone[i].isdigit():
            phone[i] = str(random.randint(0, 9))
    return ''.join(phone)

def randomize_email(email):
    if '@' in email:
        local, domain = email.split('@')
        return ''.join(random.choices(string.ascii_lowercase, k=len(local))) + '@' + domain
    return email

def anonymize_text(text):
    text = str(text)
    text = re.sub(r'\b[\w.-]+@[\w.-]+\.\w+\b', lambda m: randomize_email(m.group()), text)
    text = re.sub(r'\b\d{10,}\b', lambda m: randomize_phone(m.group()), text)
    text = re.sub(r'\b(?:Mr|Mrs|Ms|Dr)\.?\s+[A-Z][a-z]+\s+[A-Z][a-z]+\b', lambda m: randomize_string(m.group()), text)
    return text

def anonymize_dataframe(df):
    df = df.copy()
    for col in df.columns:
        if df[col].dtype == object:
            df[col] = df[col].apply(lambda x: anonymize_text(str(x)))
    return df

def load_file(filepath):
    ext = os.path.splitext(filepath)[-1].lower()
    if ext == '.csv':
        return pd.read_csv(filepath)
    elif ext == '.xlsx':
        return pd.read_excel(filepath)
    elif ext == '.txt':
        with open(filepath, 'r') as f:
            lines = f.readlines()
        return pd.DataFrame(lines, columns=['Text'])
    elif ext == '.docx':
        doc = Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return pd.DataFrame(paragraphs, columns=['Text'])
    elif ext == '.pdf':
        try:
            with pdfplumber.open(filepath) as pdf:
                text = ''.join(page.extract_text() + '\n' for page in pdf.pages if page.extract_text())
                lines = text.split('\n')
                return pd.DataFrame(lines, columns=['Text'])
        except Exception:
            with fitz.open(filepath) as doc:
                text = ''.join(page.get_text() for page in doc)
                lines = text.split('\n')
                return pd.DataFrame(lines, columns=['Text'])
    else:
        raise ValueError("Unsupported file format")

def save_anonymized(df, output_path):
    ext = os.path.splitext(output_path)[-1].lower()
    if ext == '.csv':
        df.to_csv(output_path, index=False)
    elif ext == '.xlsx':
        df.to_excel(output_path, index=False)
    elif ext == '.txt':
        with open(output_path, 'w') as f:
            for line in df.iloc[:, 0]:
                f.write(str(line) + '\n')
    elif ext == '.docx':
        doc = Document()
        for line in df.iloc[:, 0]:
            if pd.notnull(line):
                doc.add_paragraph(str(line))
        doc.save(output_path)
    else:
        raise ValueError("Unsupported save format")

# ========== API ==========

@app.route('/')
def serve_index():
    return send_from_directory('.', 'index.html')

@app.route('/anonymize', methods=['POST'])
def anonymize_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    uploaded_file = request.files['file']
    if uploaded_file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    output_format = request.form.get('output_format', '.csv').lower()
    if output_format not in ['.csv', '.xlsx', '.txt', '.docx']:
        return jsonify({"error": "Unsupported output format"}), 400

    filename = secure_filename(uploaded_file.filename)
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    uploaded_file.save(input_path)

    try:
        df = load_file(input_path)
        anon_df = anonymize_dataframe(df)

        base_name = os.path.splitext(filename)[0]
        output_filename = f'anonymized_{base_name}{output_format}'
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        save_anonymized(anon_df, output_path)
        return send_file(output_path, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ========== Main ==========

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
=======
import os
import pandas as pd
import random
import string
import re
from flask import Flask, request, send_file, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from docx import Document
import fitz  # PyMuPDF
import pdfplumber
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ========== Helper Functions ==========

def randomize_string(s):
    return ''.join(random.choice(string.ascii_letters) if c.isalpha() else c for c in s)

def randomize_phone(phone):
    phone = list(phone)
    for i in random.sample(range(min(5, len(phone))), k=3):
        if phone[i].isdigit():
            phone[i] = str(random.randint(0, 9))
    return ''.join(phone)

def randomize_email(email):
    if '@' in email:
        local, domain = email.split('@')
        return ''.join(random.choices(string.ascii_lowercase, k=len(local))) + '@' + domain
    return email

def anonymize_text(text):
    text = str(text)
    text = re.sub(r'\b[\w.-]+@[\w.-]+\.\w+\b', lambda m: randomize_email(m.group()), text)
    text = re.sub(r'\b\d{10,}\b', lambda m: randomize_phone(m.group()), text)
    text = re.sub(r'\b(?:Mr|Mrs|Ms|Dr)\.?\s+[A-Z][a-z]+\s+[A-Z][a-z]+\b', lambda m: randomize_string(m.group()), text)
    return text

def anonymize_dataframe(df):
    df = df.copy()
    for col in df.columns:
        if df[col].dtype == object:
            df[col] = df[col].apply(lambda x: anonymize_text(str(x)))
    return df

def load_file(filepath):
    ext = os.path.splitext(filepath)[-1].lower()
    if ext == '.csv':
        return pd.read_csv(filepath)
    elif ext == '.xlsx':
        return pd.read_excel(filepath)
    elif ext == '.txt':
        with open(filepath, 'r') as f:
            lines = f.readlines()
        return pd.DataFrame(lines, columns=['Text'])
    elif ext == '.docx':
        doc = Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return pd.DataFrame(paragraphs, columns=['Text'])
    elif ext == '.pdf':
        try:
            with pdfplumber.open(filepath) as pdf:
                text = ''.join(page.extract_text() + '\n' for page in pdf.pages if page.extract_text())
                lines = text.split('\n')
                return pd.DataFrame(lines, columns=['Text'])
        except Exception:
            with fitz.open(filepath) as doc:
                text = ''.join(page.get_text() for page in doc)
                lines = text.split('\n')
                return pd.DataFrame(lines, columns=['Text'])
    else:
        raise ValueError("Unsupported file format")

def save_anonymized(df, output_path):
    ext = os.path.splitext(output_path)[-1].lower()
    if ext == '.csv':
        df.to_csv(output_path, index=False)
    elif ext == '.xlsx':
        df.to_excel(output_path, index=False)
    elif ext == '.txt':
        with open(output_path, 'w') as f:
            for line in df.iloc[:, 0]:
                f.write(str(line) + '\n')
    elif ext == '.docx':
        doc = Document()
        for line in df.iloc[:, 0]:
            if pd.notnull(line):
                doc.add_paragraph(str(line))
        doc.save(output_path)
    else:
        raise ValueError("Unsupported save format")

# ========== API ==========

@app.route('/')
def serve_index():
    return send_from_directory('.', 'index.html')

@app.route('/anonymize', methods=['POST'])
def anonymize_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    uploaded_file = request.files['file']
    if uploaded_file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    output_format = request.form.get('output_format', '.csv').lower()
    if output_format not in ['.csv', '.xlsx', '.txt', '.docx']:
        return jsonify({"error": "Unsupported output format"}), 400

    filename = secure_filename(uploaded_file.filename)
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    uploaded_file.save(input_path)

    try:
        df = load_file(input_path)
        anon_df = anonymize_dataframe(df)

        base_name = os.path.splitext(filename)[0]
        output_filename = f'anonymized_{base_name}{output_format}'
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        save_anonymized(anon_df, output_path)
        return send_file(output_path, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ========== Main ==========

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
